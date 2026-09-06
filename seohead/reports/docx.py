"""Write a client-facing Word report organized as a narrative document.

Unlike the Excel workbook, this format is not a sortable data table. It presents
the conclusion first and the supporting evidence afterward. Findings therefore
appear from highest to lowest severity, while the page table is intentionally
limited and placed at the end because Word is not used for interactive sorting.
"""

from __future__ import annotations

import pathlib
from typing import Any

_MAX_PAGES_IN_TABLE = 60
_MAX_FINDINGS_PER_LEVEL = 40


def write(document: dict[str, Any], path: pathlib.Path) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    from seohead.reports import SEVERITY_TITLES

    doc = Document()
    summary = document.get("summary") or {}
    by_sev = summary.get("findings_by_severity") or {}

    doc.add_heading(f"SEO Audit: {document.get('domain', '')}", level=0)
    sub = doc.add_paragraph(document.get("url", ""))
    sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph(f"Generated: {document.get('generated_at', '')}").runs[0].font.size = Pt(9)

    # Failed/partial crawl scope must appear before the executive summary
    # table, not buried in a trailing note: a recipient must not mistake a
    # failed or sampled crawl for a clean, site-wide audit (#361).
    if summary.get("crawl_valid") is False:
        reason = summary.get("crawl_invalid_reason") or "the crawl produced no usable data"
        warn = doc.add_paragraph()
        run = warn.add_run(f"Crawl failed — no health score. {reason}")
        run.bold = True
        run.font.color.rgb = RGBColor(0xC0, 0, 0)
    if summary.get("crawl_partial"):
        finish = summary.get("crawl_finish_reason")
        scope = summary.get("crawl_scope_note")
        bits = [b for b in (f"stopped: {finish}" if finish else None, scope) if b]
        detail = f" {'; '.join(bits)}" if bits else ""
        warn = doc.add_paragraph()
        run = warn.add_run(f"Partial crawl — scope is limited.{detail}")
        run.bold = True
        run.font.color.rgb = RGBColor(0xC0, 0, 0)

    doc.add_heading("Executive Summary", level=1)
    table = doc.add_table(rows=0, cols=2)
    table.style = "Light Grid Accent 1"
    for name, value in (
        ("Pages checked", summary.get("pages_checked", 0)),
        ("Critical issues", by_sev.get("critical", 0)),
        ("Warnings", by_sev.get("warning", 0)),
        ("Notices", by_sev.get("notice", 0)),
    ):
        row = table.add_row().cells
        row[0].text, row[1].text = name, str(value)

    disabled = summary.get("checks_disabled") or []
    if disabled:
        doc.add_heading("Disabled Checks", level=1)
        doc.add_paragraph(
            "These checks were deliberately turned off and did not run. "
            "Their silence is a configuration choice, not a clean result:"
        )
        for item in disabled:
            doc.add_paragraph(f"{item.get('id')} — {item.get('reason')}", style="List Bullet")

    failed = summary.get("tools_failed") or []
    if failed:
        doc.add_heading("Unavailable Checks", level=1)
        doc.add_paragraph(
            "These checks did not complete, so their evidence is absent from the report. "
            "Their silence does not mean that no issues were found:"
        )
        for item in failed:
            doc.add_paragraph(f"{item.get('tool')} — {item.get('error')}", style="List Bullet")

    findings = document.get("findings") or []
    for level in ("critical", "warning", "notice"):
        chunk = [f for f in findings if f.get("severity") == level]
        if not chunk:
            continue
        doc.add_heading(f"{SEVERITY_TITLES.get(level, level)} — {len(chunk)}", level=1)
        if level == "critical":
            warn = doc.add_paragraph()
            run = warn.add_run(
                "These are the highest-severity issues this audit found, by the "
                "aggregator's severity rules. This report does not measure current "
                "search rankings, so it does not say whether these issues have "
                "already cost the site any."
            )
            run.bold = True
            run.font.color.rgb = RGBColor(0xC0, 0, 0)
        for finding in chunk[:_MAX_FINDINGS_PER_LEVEL]:
            text = finding.get("text", "")
            where = finding.get("url") or finding.get("source", "")
            para = doc.add_paragraph(style="List Bullet")
            para.add_run(text)
            if where:
                tail = para.add_run(f"  [{where}]")
                tail.font.size = Pt(8)
                tail.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        if len(chunk) > _MAX_FINDINGS_PER_LEVEL:
            doc.add_paragraph(
                f"…and {len(chunk) - _MAX_FINDINGS_PER_LEVEL} more. "
                "See the Excel version of this audit for the complete list."
            )

    pages = document.get("pages") or []
    if pages:
        doc.add_heading("Pages", level=1)
        table = doc.add_table(rows=1, cols=5)
        table.style = "Light Grid Accent 1"
        for i, title in enumerate(("URL", "Status", "Title", "Words", "Canonical")):
            table.rows[0].cells[i].text = title
        for page in pages[:_MAX_PAGES_IN_TABLE]:
            cells = table.add_row().cells
            for index, name, limit in (
                (0, "url", None),
                (1, "status", None),
                (2, "title", 120),
                (3, "words", None),
                (4, "canonical", 120),
            ):
                value = page.get(name)
                text = "" if value is None else str(value)
                cells[index].text = text[:limit] if limit else text
        if len(pages) > _MAX_PAGES_IN_TABLE:
            doc.add_paragraph(f"Showing the first {_MAX_PAGES_IN_TABLE} of {len(pages)} pages.")

    note = summary.get("severity_note")
    if note:
        doc.add_paragraph()
        tail = doc.add_paragraph(note)
        tail.runs[0].font.size = Pt(8)
        tail.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.save(str(path))
