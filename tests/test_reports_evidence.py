"""Regressions for #337, #338, #361: report-build schema recognition and evidence.

All fixtures are synthetic and offline; no network access, no real credentials.
"""

from __future__ import annotations

import copy
import csv
import json

from openpyxl import load_workbook

from seohead.audit.site import SCHEMA
from seohead.reports import build_report

# ── #338: report-build must recognize the declared audit schema version ────

SITE_DOCUMENT = {
    "schema": SCHEMA,
    "domain": "example.com",
    "url": "https://example.com/",
    "generated_at": "2026-09-05T00:00:00Z",
    "findings": [],
    "pages": [],
    "summary": {"pages_checked": 0, "findings_total": 0, "findings_by_severity": {}},
}

SF_DOCUMENT = {
    "schema_version": "2.0",
    "run": {"project": "example.test", "source": "https://example.test/"},
    "summary": {
        "totals": {"urls_crawled": 1, "issues_total": 0},
        "by_severity": {"critical": 0, "warning": 0, "notice": 0},
        "by_check": {},
    },
    "issues": [],
    "pages": [{"url": "https://example.test/", "status_code": 200, "metrics": {}}],
    "groups": [],
}


def test_unsupported_site_audit_schema_is_refused_and_names_the_marker(tmp_path):
    doc = {**SITE_DOCUMENT, "schema": "seohead.site-audit/999"}
    result = build_report(doc, fmt="md", path=str(tmp_path / "out.md"))
    assert result["ok"] is False
    assert "schema" in result["error"]
    assert "seohead.site-audit/999" in result["error"]


def test_missing_site_audit_schema_is_refused():
    doc = {k: v for k, v in SITE_DOCUMENT.items() if k != "schema"}
    result = build_report(doc, fmt="md")
    assert result["ok"] is False
    assert "schema" in result["error"]


def test_marked_site_audit_with_malformed_required_containers_is_refused(tmp_path):
    """#338: the marker alone is not enough for a writer to safely consume the document."""
    for field, value in (("findings", "not a list"), ("pages", {}), ("summary", [])):
        doc = copy.deepcopy(SITE_DOCUMENT)
        doc[field] = value
        result = build_report(doc, fmt="md", path=str(tmp_path / f"site-{field}.md"))
        assert result["ok"] is False
        assert field in result["error"]


def test_unsupported_sf_schema_version_is_refused_and_names_the_marker(tmp_path):
    doc = {**SF_DOCUMENT, "schema_version": "999.0"}
    result = build_report(doc, fmt="md", path=str(tmp_path / "out.md"))
    assert result["ok"] is False
    assert "schema_version" in result["error"]
    assert "999.0" in result["error"]


def test_marked_sf_audit_with_malformed_required_containers_is_refused(tmp_path):
    """#338: reject every container normalization or a writer would dereference."""
    for field, value in (
        ("run", []),
        ("summary", "not a mapping"),
        ("issues", {}),
        ("pages", {}),
        ("groups", {}),
    ):
        doc = copy.deepcopy(SF_DOCUMENT)
        doc[field] = value
        result = build_report(doc, fmt="md", path=str(tmp_path / f"sf-{field}.md"))
        assert result["ok"] is False
        assert field in result["error"]


def test_correctly_versioned_site_audit_still_renders(tmp_path):
    """Positive control: the exact supported marker must not be caught by the gate."""
    result = build_report(SITE_DOCUMENT, fmt="md", path=str(tmp_path / "out.md"))
    assert result["ok"] is True, result


def test_correctly_versioned_sf_audit_still_renders(tmp_path):
    """Positive control: the exact supported marker must not be caught by the gate."""
    result = build_report(SF_DOCUMENT, fmt="md", path=str(tmp_path / "out.md"))
    assert result["ok"] is True, result


def test_json_format_is_also_gated_on_an_unsupported_marker():
    """#338: the recognition gate applies to json too, only the copy step is untouched."""
    doc = {**SF_DOCUMENT, "schema_version": "999.0"}
    result = build_report(doc, fmt="json")
    assert result["ok"] is False
    assert "schema_version" in result["error"]


def test_json_output_of_a_recognized_document_is_an_untouched_copy(tmp_path):
    target = tmp_path / "out.json"
    result = build_report(SF_DOCUMENT, fmt="json", path=str(target))
    assert result["ok"] is True
    assert json.loads(target.read_text(encoding="utf-8")) == SF_DOCUMENT


# ── #337: "Checks completed" must not be a finding-only list's length ──────


def test_clean_sf_audit_reports_checks_completed_as_unavailable_not_zero(tmp_path):
    """A clean run (0 findings) must not render as if 0 checks ran (#337)."""
    target = tmp_path / "clean.xlsx"
    result = build_report(SF_DOCUMENT, fmt="xlsx", path=str(target))
    assert result["ok"], result
    ws = load_workbook(target)["Summary"]
    rows = {row[0]: row[1] for row in ws.iter_rows(values_only=True) if row[0]}
    assert rows["Checks completed"] == "not reported"


def test_sf_audit_with_check_coverage_reports_a_real_completed_count(tmp_path):
    """When the source supplies check_coverage, the true completed inventory is shown."""
    doc = copy.deepcopy(SF_DOCUMENT)
    doc["summary"]["by_check"] = {"BROKEN_INTERNAL_LINK": 3}
    doc["summary"]["check_coverage"] = {
        "checks_total": 5,
        "checks_fired": 1,
        "checks_skipped": 0,
        "checks_disabled": 0,
        "checks_disabled_ids": [],
        "checks_silent": 3,
        "checks_silent_ids": ["TITLE_MISSING", "H1_MISSING", "CANONICAL_MISSING"],
        "coverage": 1.0,
    }
    target = tmp_path / "covered.xlsx"
    result = build_report(doc, fmt="xlsx", path=str(target))
    assert result["ok"], result
    ws = load_workbook(target)["Summary"]
    rows = {row[0]: row[1] for row in ws.iter_rows(values_only=True) if row[0]}
    # fired (1) + silent (3) = 4 checks actually completed -- not len(by_check) == 1.
    assert rows["Checks completed"] == 4


def test_site_audit_contract_keeps_its_own_real_tools_run_count(tmp_path):
    """Negative control: the site-audit contract's own tools_run must render as before."""
    doc = copy.deepcopy(SITE_DOCUMENT)
    doc["summary"]["tools_run"] = ["cdn_check", "render_check"]
    target = tmp_path / "site.xlsx"
    result = build_report(doc, fmt="xlsx", path=str(target))
    assert result["ok"], result
    ws = load_workbook(target)["Summary"]
    rows = {row[0]: row[1] for row in ws.iter_rows(values_only=True) if row[0]}
    assert rows["Checks completed"] == 2


# ── #361: failed/partial scope and disabled-check evidence must survive ────

_PARTIAL_SF_AUDIT = {
    "schema_version": "2.0",
    "run": {
        "project": "example.test",
        "source": "https://example.test/",
        "crawl_valid": False,
        "crawl_invalid_reason": "no response",
        "crawl_partial": True,
        "crawl_finish_reason": "url_limit",
        "crawl_stopped_reason": "url_limit",
        "checks_skipped": [],
        "checks_disabled": [],
    },
    "summary": {
        "totals": {"urls_crawled": 1, "issues_total": 1},
        "by_severity": {"critical": 1, "warning": 0, "notice": 0},
        "health_score": None,
        "health_score_reason": "no response",
        "health_score_scope": "1 of 1,000 sitemap URLs crawled",
    },
    "issues": [
        {
            "check": "NO_RESPONSE",
            "severity": "critical",
            "message": "No response",
            "target_url": "https://example.test/",
        }
    ],
    "pages": [{"url": "https://example.test/", "metrics": {}}],
    "groups": [],
}


def test_invalid_partial_sf_audit_surfaces_scope_evidence_in_md(tmp_path):
    target = tmp_path / "partial.md"
    result = build_report(_PARTIAL_SF_AUDIT, fmt="md", path=str(target))
    assert result["ok"], result
    text = target.read_text(encoding="utf-8")
    assert "Crawl failed" in text and "no response" in text
    assert "Partial crawl" in text and "url_limit" in text
    assert "1 of 1,000 sitemap URLs crawled" in text
    # Scope evidence must precede the metrics table (#361).
    assert text.index("Crawl failed") < text.index("| Pages checked |")


def test_invalid_partial_sf_audit_surfaces_scope_evidence_in_xlsx(tmp_path):
    target = tmp_path / "partial.xlsx"
    result = build_report(_PARTIAL_SF_AUDIT, fmt="xlsx", path=str(target))
    assert result["ok"], result
    ws = load_workbook(target)["Summary"]
    joined = "\n".join(str(c.value or "") for row in ws.iter_rows() for c in row)
    assert "Crawl failed" in joined and "no response" in joined
    assert "url_limit" in joined
    assert "1 of 1,000 sitemap URLs crawled" in joined


def test_invalid_partial_sf_audit_surfaces_scope_evidence_in_docx(tmp_path):
    from docx import Document

    target = tmp_path / "partial.docx"
    result = build_report(_PARTIAL_SF_AUDIT, fmt="docx", path=str(target))
    assert result["ok"], result
    text = "\n".join(p.text for p in Document(str(target)).paragraphs)
    assert "Crawl failed" in text and "no response" in text
    assert "Partial crawl" in text and "url_limit" in text


def test_disabled_checks_are_distinct_from_skipped_checks_in_md(tmp_path):
    doc = copy.deepcopy(_PARTIAL_SF_AUDIT)
    doc["run"]["crawl_valid"] = True
    doc["run"]["crawl_partial"] = False
    doc["run"]["checks_skipped"] = [{"id": "SF_LOG_ANALYZE", "reason": "log file unavailable"}]
    doc["run"]["checks_disabled"] = [{"id": "BROKEN_PAGE_4XX", "reason": "disabled in config"}]
    doc["issues"] = []
    target = tmp_path / "mixed.md"
    result = build_report(doc, fmt="md", path=str(target))
    assert result["ok"], result
    text = target.read_text(encoding="utf-8")
    assert "## Disabled checks" in text
    assert "BROKEN_PAGE_4XX" in text and "disabled in config" in text
    assert "## Unavailable checks" in text
    assert "SF_LOG_ANALYZE" in text and "log file unavailable" in text
    # The two sections must not merge into one list.
    assert text.index("## Disabled checks") != text.index("## Unavailable checks")


def test_complete_sf_audit_has_no_scope_warning(tmp_path):
    """Negative control: a complete, non-disabled audit renders no warning at all (#361)."""
    doc = copy.deepcopy(_PARTIAL_SF_AUDIT)
    doc["run"]["crawl_valid"] = True
    doc["run"]["crawl_partial"] = False
    doc["run"]["checks_disabled"] = []
    doc["issues"] = []
    for target, fmt in (
        (tmp_path / "clean.md", "md"),
        (tmp_path / "clean.xlsx", "xlsx"),
    ):
        result = build_report(doc, fmt=fmt, path=str(target))
        assert result["ok"], result
    md_text = (tmp_path / "clean.md").read_text(encoding="utf-8")
    assert "Crawl failed" not in md_text
    assert "Partial crawl" not in md_text
    assert "Disabled checks" not in md_text
    ws = load_workbook(tmp_path / "clean.xlsx")["Summary"]
    joined = "\n".join(str(c.value or "") for row in ws.iter_rows() for c in row)
    assert "Crawl failed" not in joined
    assert "Partial crawl" not in joined
    assert "Disabled check" not in joined


# ── #574: CSV must retain the same run-scope evidence as other reports ─────


def test_csv_writes_scope_evidence_without_polluting_tracker_findings(tmp_path):
    """A failed, partial run cannot be indistinguishable from a clean zero-issue CSV."""
    target = tmp_path / "partial.csv"
    doc = copy.deepcopy(_PARTIAL_SF_AUDIT)
    doc["issues"] = []
    doc["pages"] = []
    doc["run"]["checks_disabled"] = [{"id": "BROKEN_PAGE_4XX", "reason": "disabled in config"}]
    doc["run"]["checks_skipped"] = [{"id": "SF_LOG_ANALYZE", "reason": "log file unavailable"}]
    from seohead.servers.handlers import report_build

    result = report_build(audit=doc, fmt="csv", out=str(target))
    assert result["ok"], result
    assert result["outputs"] == [
        str(target),
        str(target.with_suffix(".scope.csv")),
    ]

    with target.open(encoding="utf-8-sig", newline="") as fh:
        findings = list(csv.reader(fh, delimiter=";"))
    assert len(findings) == 1  # Scope evidence cannot be mistaken for a tracker finding.

    with target.with_suffix(".scope.csv").open(encoding="utf-8-sig", newline="") as fh:
        scope = list(csv.DictReader(fh, delimiter=";"))
    assert scope == [
        {
            "Evidence type": "crawl",
            "Identifier": "validity",
            "Status": "failed",
            "Reason": "no response",
        },
        {
            "Evidence type": "crawl",
            "Identifier": "scope",
            "Status": "partial",
            "Reason": "stopped: url_limit; 1 of 1,000 sitemap URLs crawled",
        },
        {
            "Evidence type": "check",
            "Identifier": "BROKEN_PAGE_4XX",
            "Status": "disabled",
            "Reason": "disabled in config",
        },
        {
            "Evidence type": "check",
            "Identifier": "SF_LOG_ANALYZE",
            "Status": "unavailable",
            "Reason": "log file unavailable",
        },
    ]


def test_clean_csv_has_an_empty_scope_sidecar(tmp_path):
    """A fresh clean report cannot retain caveats from an earlier write to the same path."""
    target = tmp_path / "clean.csv"
    result = build_report(SF_DOCUMENT, fmt="csv", path=str(target))
    assert result["ok"], result
    with target.with_suffix(".scope.csv").open(encoding="utf-8-sig", newline="") as fh:
        assert list(csv.reader(fh, delimiter=";")) == [
            ["Evidence type", "Identifier", "Status", "Reason"]
        ]


# ── #575: present-but-None page facts are absent, not the text "None" ────


def test_docx_and_markdown_leave_none_page_fields_blank(tmp_path):
    doc = copy.deepcopy(SITE_DOCUMENT)
    doc["pages"] = [
        {
            "url": "https://example.com/unavailable",
            "status": None,
            "title": None,
            "words": None,
            "canonical": None,
        }
    ]

    md_target = tmp_path / "none.md"
    docx_target = tmp_path / "none.docx"
    assert build_report(doc, fmt="md", path=str(md_target))["ok"]
    assert build_report(doc, fmt="docx", path=str(docx_target))["ok"]

    page_row = next(
        line for line in md_target.read_text(encoding="utf-8").splitlines() if "unavailable" in line
    )
    assert page_row == "| https://example.com/unavailable |  |  |  |  |"

    from docx import Document

    document = Document(str(docx_target))
    page_table = next(table for table in document.tables if table.cell(0, 0).text == "URL")
    assert [cell.text for cell in page_table.rows[1].cells] == [
        "https://example.com/unavailable",
        "",
        "",
        "",
        "",
    ]
