"""Every command shown in the documentation must still run.

``scripts/doc_commands.py`` extracts every ``seohead ...`` invocation from the public
Markdown (README, docs/, skills, examples). This module turns each one into something
runnable entirely offline — URLs point at a loopback fixture server instead of the live
internet, and file/directory arguments point at materialized copies of ``examples/`` —
then either executes it in-process through :func:`seohead.cli.main` (asserting a clean
exit) or, for the handful that fundamentally need infrastructure no local fixture can
stand in for (a real RDAP/DNS ecosystem, a licensed Screaming Frog binary, a paid
provider credential, a server that never returns), parses it with the real argument
parser so a renamed or removed flag still fails the build.

A documented command that no longer works is worse than a missing one, because it is
trusted (issue #22).
"""

from __future__ import annotations

import contextlib
import io
import json
import shutil
from pathlib import Path

import pytest

from scripts.doc_commands import extract_commands, to_argv
from tests.doc_fixtures.site_server import run_fixture_site

ROOT = Path(__file__).resolve().parent.parent

# Each needs something a local, credential-free fixture cannot stand in for:
# a real RDAP/WHOIS/DNS ecosystem, a licensed Screaming Frog binary or a captured
# .seospider file, a paid or credentialed external API, or (for `mcp`) a server
# that never returns. Full execution is skipped for these; CLI parseability
# (the flags still exist) is still checked below.
NEEDS_LIVE_INFRASTRUCTURE = {
    "domain-profile",
    "mirror-check",
    # Needs a headless browser. Playwright is an optional extra, so a runner without it gets a
    # truthful ok:false ("playwright is not installed") — which, since #155 made ok:false reach
    # the exit code, is now a non-zero exit rather than a silently swallowed success. That is the
    # command working, not failing: the same distinction sf doctor and save-config already draw
    # between "this environment lacks the dependency" and "this command is broken".
    "render-check",
    # Its target-domain matching deliberately rejects a bare IP address (the fixture server has
    # no name, only 127.0.0.1) — the same "needs a real domain, not just an HTTP-answering host"
    # gap as domain-profile and mirror-check above, just surfaced only once ok:false reached the
    # exit code (#155) instead of being silently swallowed by an exit-0 success.
    "backlinks-check",
    "keywords-expand",
    "keywords-seasonality",
    "keywords-exact",
    "serp-fetch",
    "google-keywords",
    "google-serp",
    "metrika-counters",
    "metrika-setup",
    "metrika-report",
    "regions-tree",
    "mcp",
}


# `sf` subcommands that inspect the local Screaming Frog installation itself rather than an
# export: they pass on a developer machine that has SF and fail on a runner that does not,
# which is the difference between an environment and a broken command.
SF_SUBCOMMANDS_NEEDING_AN_INSTALL = {"doctor", "save-config"}


def _is_licensed_sf_mode(argv: list[str]) -> bool:
    """Mode A (`--crawl`) needs the licensed SF CLI; `--load-crawl` needs a real capture;
    `doctor` and `save-config` need an installed Screaming Frog to report on at all."""
    if argv[:1] != ["sf"]:
        return False
    if any(flag in argv for flag in ("--crawl", "--load-crawl")):
        return True
    return argv[1:2] and argv[1] in SF_SUBCOMMANDS_NEEDING_AN_INSTALL


def _substitute(raw: str, base_url: str) -> str:
    """Point every URL/domain placeholder in a documented command at the fixture site."""
    host = base_url.split("//", 1)[1]
    # Longer / prefixed patterns first, so a bare fallback cannot half-rewrite one.
    replacements = [
        ("https://example-msk.example", base_url),
        ("https://example-spb.example", base_url),
        ("https://donor1.example", base_url),
        ("https://donor2.example", base_url),
        ("https://<domain>", base_url),
        ("https://example.org", base_url),
        ("https://example.com", base_url),
        ("<domain>", host),
        ("example.com", host),
        ("<template>", "page"),
        ("./sf-exports", "exports"),
        ("./exports", "exports"),
    ]
    text = raw
    for old, new in replacements:
        text = text.replace(old, new)
    return text


def _seed_workdir(tmp_path: Path, base_url: str) -> None:
    """Materialize every fixture a documented command's relative path expects."""
    shutil.copytree(ROOT / "examples", tmp_path / "examples")
    shutil.copytree(ROOT / "examples" / "exports", tmp_path / "exports")
    shutil.copy(ROOT / "examples" / "audit.json", tmp_path / "audit.json")
    shutil.copy(ROOT / "examples" / "audit.json", tmp_path / "old-audit.json")
    shutil.copy(ROOT / "examples" / "audit.json", tmp_path / "new-audit.json")
    shutil.copy(ROOT / "config.example.json", tmp_path / "config.json")
    # A finished, internally consistent crawl output, so a documented `log-scan --run ./run`
    # actually scans something instead of reporting that the directory is empty.
    shutil.copytree(ROOT / "tests" / "doc_fixtures" / "run", tmp_path / "run")
    (tmp_path / "report").mkdir()
    shutil.copy(ROOT / "examples" / "audit.json", tmp_path / "report" / "audit.json")
    (tmp_path / "crawl.json").write_text(json.dumps({"limits": {"max_urls": 5}}), encoding="utf-8")
    (tmp_path / "donors.txt").write_text(f"{base_url}/page\n", encoding="utf-8")
    images_dir = tmp_path / "images"
    images_dir.mkdir()
    shutil.copy(ROOT / "tests" / "doc_fixtures" / "site" / "image.png", images_dir / "image.png")
    # docs/../deliverables.md's chain optimizes whatever `images-download --output-dir
    # ./original` just wrote; its own `--urls` is an illustrative "<comma list>" placeholder
    # that downloads nothing, so pre-seed the directory the second command reads from.
    original_dir = tmp_path / "original"
    original_dir.mkdir()
    shutil.copy(ROOT / "tests" / "doc_fixtures" / "site" / "image.png", original_dir / "image.png")


def _seed_scan_inputs(tmp_path: Path) -> None:
    from dataclasses import asdict

    from seohead.crawl.collect import PageRecord
    from seohead.storage import import_run

    source = tmp_path / "scan-inputs"
    source.mkdir()
    audit = json.loads((tmp_path / "audit.json").read_text(encoding="utf-8"))
    (source / "audit.json").write_bytes((tmp_path / "audit.json").read_bytes())
    (source / "pages.jsonl").write_text(
        "".join(
            json.dumps(asdict(PageRecord(url=page["url"], status_code=page.get("status_code"))))
            + "\n"
            for page in audit["pages"]
        ),
        encoding="utf-8",
    )
    (source / "links.jsonl").write_text("", encoding="utf-8")
    original_config = audit["run"].get("crawl_config") or {"limits.max_urls": len(audit["pages"])}
    scan = import_run(
        source, tmp_path / "scan.sqlite", producer_build="1" * 40, effective_config=original_config
    )
    for name in ("before.sqlite", "after.sqlite"):
        shutil.copyfile(scan, tmp_path / name)


@pytest.fixture(scope="module")
def fixture_site():
    with run_fixture_site() as base_url:
        yield base_url


def _cases():
    for command in extract_commands(ROOT):
        argv = to_argv(command.raw)
        if not argv:
            continue
        yield pytest.param(command, argv, id=f"{command.source.name}:{command.raw}"[:120])


@pytest.mark.parametrize("command,argv", list(_cases()))
def test_documented_command_executes_or_at_least_still_parses(
    command, argv, fixture_site, tmp_path, monkeypatch
):
    tool = argv[0]
    if tool in NEEDS_LIVE_INFRASTRUCTURE or _is_licensed_sf_mode(argv):
        from seohead.cli import build_parser
        from seohead.sf.cli import build_parser as build_sf_parser

        try:
            if argv[0] == "sf":
                build_sf_parser().parse_args(argv[1:])
            else:
                build_parser().parse_args(argv)
        except SystemExit as exc:
            assert exc.code == 0, (
                f"{command.source.relative_to(ROOT)}: `{command.raw}` no longer parses"
            )
        return

    from seohead.cli import main as cli_main

    _seed_workdir(tmp_path, fixture_site)
    if any(".sqlite" in value for value in argv):
        _seed_scan_inputs(tmp_path)
    if command.source.name == "robots-blocked.md":
        (tmp_path / "config.json").write_text(
            json.dumps({"robots": {"user_agent_token": "Googlebot"}}), encoding="utf-8"
        )
    monkeypatch.chdir(tmp_path)
    monkeypatch.setenv("SEOHEAD_ALLOW_PRIVATE_NETWORKS", "1")
    # A command with no explicit `echo ... |` payload still probes stdin for JSON
    # input; pytest's own captured stdin raises on read instead of giving EOF, so
    # every case (not only the piped ones) gets a real, harmless stream here.
    monkeypatch.setattr("sys.stdin", io.StringIO(command.stdin or ""))

    substituted = to_argv(_substitute(command.raw, fixture_site))
    if substituted[:1] == ["site-audit"] and "--skip" not in substituted:
        # site-audit's own domain-profile sub-check needs real RDAP; every other
        # site-level tool in it is a plain HTTP GET the fixture server answers.
        substituted += ["--skip", "domain_profile"]

    stdout = io.StringIO()
    try:
        with contextlib.redirect_stdout(stdout):
            exit_code = cli_main(substituted)
    except SystemExit as exc:  # --version / --help exit through argparse itself
        exit_code = exc.code if isinstance(exc.code, int) else 1

    assert exit_code == 0, (
        f"{command.source.relative_to(ROOT)}: `{command.raw}` exited {exit_code}: "
        f"{stdout.getvalue()[-2000:]}"
    )


def test_every_documented_command_was_exercised_above():
    """Guards the extractor itself: a doc command that fails to parse must fail loudly,
    not vanish from the parametrized list above."""
    commands = extract_commands(ROOT)
    assert commands, "no documented `seohead` commands were found at all"
    for command in commands:
        to_argv(command.raw)  # raises on the rare unparseable line


def test_sf_run_then_report_build_renders_the_real_totals(tmp_path, monkeypatch):
    """The exact two-line recipe docs/USAGE.md and docs/TOOLS.md document, checked on
    *content* rather than exit code alone (#151): `sf run` writes findings under
    ``issues``/``pages[].metrics``, and a renderer that only reads ``findings``/flat
    page keys silently produces a confident 0/0/0/0 report for a site that failed
    its audit. The parametrized test above already runs this recipe and only checks
    that it exits 0 -- which a report full of zeros still does.
    """
    from seohead.cli import main as cli_main

    monkeypatch.chdir(tmp_path)
    shutil.copytree(ROOT / "examples", tmp_path / "examples")

    assert (
        cli_main(["sf", "run", "--exports-dir", "examples/exports", "--out", "report", "--tasks"])
        == 0
    )
    audit = json.loads((tmp_path / "report" / "audit.json").read_text(encoding="utf-8"))
    totals = audit["summary"]["totals"]
    by_severity = audit["summary"]["by_severity"]
    assert by_severity["critical"] > 0, "fixture must exercise a report with real findings"

    docx_path = tmp_path / "client.docx"
    assert (
        cli_main(
            [
                "report-build",
                "--audit",
                "report/audit.json",
                "--format",
                "docx",
                "--out",
                str(docx_path),
            ]
        )
        == 0
    )
    from docx import Document

    doc = Document(str(docx_path))
    summary_row = {row.cells[0].text: row.cells[1].text for row in doc.tables[0].rows}
    assert summary_row["Pages checked"] == str(totals["urls_crawled"])
    assert summary_row["Critical issues"] == str(by_severity["critical"])
    assert summary_row["Warnings"] == str(by_severity["warning"])
    assert summary_row["Notices"] == str(by_severity["notice"])
    # Not just the summary table: the per-severity findings sections themselves,
    # the part a zero-findings render dropped entirely.
    paragraph_text = "\n".join(p.text for p in doc.paragraphs)
    assert f"Critical — {by_severity['critical']}" in paragraph_text
    first_issue = next(i for i in audit["issues"] if i["severity"] == "critical")
    assert first_issue["message"] in paragraph_text

    xlsx_path = tmp_path / "client.xlsx"
    assert (
        cli_main(
            [
                "report-build",
                "--audit",
                "report/audit.json",
                "--format",
                "xlsx",
                "--out",
                str(xlsx_path),
            ]
        )
        == 0
    )
    from openpyxl import load_workbook

    wb = load_workbook(xlsx_path)
    ws = wb["Summary"]
    metrics = {ws.cell(row=r, column=1).value: ws.cell(row=r, column=2).value for r in range(6, 11)}
    assert metrics["Pages checked"] == totals["urls_crawled"]
    assert metrics["Total findings"] == totals["issues_total"]
    assert metrics["Critical findings"] == by_severity["critical"]
    assert metrics["Warnings"] == by_severity["warning"]
    assert metrics["Notices"] == by_severity["notice"]
    assert wb["Findings"].max_row - 1 == totals["issues_total"]
    assert wb["Pages"].max_row - 1 == totals["urls_crawled"]


def test_report_build_refuses_a_document_matching_neither_schema(tmp_path):
    """Neither ``findings`` (site-audit) nor ``issues``+``pages`` (SF Analyzer): this is
    not a document either renderer contract in this package produces, so it must be
    refused loudly rather than rendered as an empty, confident report (#151).
    """
    from seohead.reports import build_report

    result = build_report({"totally": "unrelated"}, fmt="docx", path=str(tmp_path / "out.docx"))
    assert result["ok"] is False
    assert "not recognized" in result["error"]
    assert not (tmp_path / "out.docx").exists()
